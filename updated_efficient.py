#type: ignore
import os
import numpy as np
import tensorflow as tf
import matplotlib.pyplot as plt
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications.efficientnet_v2 import EfficientNetV2B3, preprocess_input
from tensorflow.keras.layers import Dense, GlobalAveragePooling2D, Dropout
from tensorflow.keras.models import Model
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.callbacks import ModelCheckpoint, EarlyStopping, ReduceLROnPlateau
from sklearn.metrics import classification_report, confusion_matrix
from sklearn.utils.class_weight import compute_class_weight
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Paths
input_path = r"D:\Research\Jujube\boroi"
output_path = r"D:\Research\Jujube\Results\Updated Efficient"
os.makedirs(output_path, exist_ok=True)

# Parameters
IMG_SIZE = (300, 300)
BATCH_SIZE = 32
EPOCHS = 30  # 🔥 Updated for longer training
LEARNING_RATE = 0.0001

# Data Augmentation and Preprocessing
train_datagen = ImageDataGenerator(
    preprocessing_function=preprocess_input,
    validation_split=0.2,
    rotation_range=40,  # 🔥 More rotation
    width_shift_range=0.2,
    height_shift_range=0.2,
    shear_range=0.2,
    zoom_range=0.3,  # 🔥 More zoom
    horizontal_flip=True,
    fill_mode="nearest"
)

# Generators
train_generator = train_datagen.flow_from_directory(
    input_path,
    target_size=IMG_SIZE,
    batch_size=BATCH_SIZE,
    class_mode='categorical',
    subset='training',
    shuffle=True
)

validation_generator = train_datagen.flow_from_directory(
    input_path,
    target_size=IMG_SIZE,
    batch_size=BATCH_SIZE,
    class_mode='categorical',
    subset='validation',
    shuffle=False
)

# 🔥 Compute Class Weights
classes = list(train_generator.class_indices.keys())
labels = train_generator.classes
class_weights = compute_class_weight(
    class_weight='balanced',
    classes=np.unique(labels),
    y=labels
)
class_weights = dict(enumerate(class_weights))

# Load EfficientNetV2
base_model = EfficientNetV2B3(weights='imagenet', include_top=False, input_shape=(300, 300, 3))
base_model.trainable = False

# Custom Top Layers
x = base_model.output
x = GlobalAveragePooling2D()(x)
x = Dense(1024, activation='relu')(x)
x = Dropout(0.5)(x)
predictions = Dense(train_generator.num_classes, activation='softmax')(x)

model = Model(inputs=base_model.input, outputs=predictions)

# Compile
model.compile(optimizer=Adam(learning_rate=LEARNING_RATE),  # 🔥 Use Adam
              loss='categorical_crossentropy',
              metrics=['accuracy'])

# Callbacks
checkpoint = ModelCheckpoint(os.path.join(output_path, 'best_model.keras'),
                             monitor='val_loss', save_best_only=True, mode='min', verbose=1)

early_stop = EarlyStopping(monitor='val_loss', patience=6, restore_best_weights=True)  # 🔥 Better stopping

reduce_lr = ReduceLROnPlateau(monitor='val_loss', factor=0.5, patience=3, verbose=1)

# Initial Training
history = model.fit(
    train_generator,
    epochs=EPOCHS,
    validation_data=validation_generator,
    class_weight=class_weights,  # 🔥 Added class weight
    callbacks=[checkpoint, early_stop, reduce_lr]
)

# Fine-Tuning (Unfreeze last few layers)
base_model.trainable = True
for layer in base_model.layers[:-30]:  # 🔥 Only last 30 layers trainable
    layer.trainable = False

# Recompile
model.compile(optimizer=Adam(learning_rate=LEARNING_RATE/10),  # Smaller LR for fine-tuning
              loss='categorical_crossentropy',
              metrics=['accuracy'])

# Fine-Tune Training
fine_tune_epochs = 10
model.fit(
    train_generator,
    epochs=fine_tune_epochs,
    validation_data=validation_generator,
    class_weight=class_weights,
    callbacks=[checkpoint, early_stop, reduce_lr]
)

# Plot Accuracy & Loss
plt.figure(figsize=(12, 5))

plt.subplot(1, 2, 1)
plt.plot(history.history['accuracy'], label='Train Accuracy', marker='o')
plt.plot(history.history['val_accuracy'], label='Validation Accuracy', marker='o')
plt.title("Model Accuracy")
plt.xlabel("Epochs")
plt.ylabel("Accuracy")
plt.legend()

plt.subplot(1, 2, 2)
plt.plot(history.history['loss'], label='Train Loss', marker='o')
plt.plot(history.history['val_loss'], label='Validation Loss', marker='o')
plt.title("Model Loss")
plt.xlabel("Epochs")
plt.ylabel("Loss")
plt.legend()

accuracy_loss_curve_path = os.path.join(output_path, 'accuracy_loss_curve.png')
plt.savefig(accuracy_loss_curve_path)
plt.close()

# Load Best Model
model.load_weights(os.path.join(output_path, 'best_model.keras'))

# Evaluation
eval_result = model.evaluate(validation_generator)

Y_pred = model.predict(validation_generator)
y_pred = np.argmax(Y_pred, axis=1)
y_true = validation_generator.classes

report = classification_report(y_true, y_pred, target_names=classes)

# Confusion Matrix
cm = confusion_matrix(y_true, y_pred)
confusion_matrix_path = os.path.join(output_path, 'confusion_matrix.png')

plt.figure(figsize=(12, 10))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
            xticklabels=classes,
            yticklabels=classes)
plt.xticks(rotation=45)
plt.yticks(rotation=0)
plt.xlabel('Predicted')
plt.ylabel('Actual')
plt.title('Confusion Matrix')
plt.savefig(confusion_matrix_path)
plt.close()

# Word Document
doc = Document()
doc.add_heading('Classification Report', 0)

# Model Performance
doc.add_heading('Model Performance', level=1)
doc.add_paragraph(f"Final Validation Loss: {eval_result[0]:.4f}")
doc.add_paragraph(f"Final Validation Accuracy: {eval_result[1]:.4f}")

# Metrics
doc.add_heading('Classification Metrics', level=1)
doc.add_paragraph(report)

# History
doc.add_heading('Training History', level=1)
doc.add_picture(accuracy_loss_curve_path, width=Inches(6))

# Confusion Matrix
doc.add_heading('Confusion Matrix', level=1)
doc.add_picture(confusion_matrix_path, width=Inches(6))

# Save
doc.save(os.path.join(output_path, 'classification_report.docx'))

print("Training completed and improved results saved!")