import os
import numpy as np
import tensorflow as tf
import matplotlib.pyplot as plt
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications.mobilenet_v2 import MobileNetV2, preprocess_input
from tensorflow.keras.layers import Dense, GlobalAveragePooling2D, Dropout
from tensorflow.keras.models import Model
from tensorflow.keras.optimizers import RMSprop
from tensorflow.keras.callbacks import ModelCheckpoint, EarlyStopping
from sklearn.metrics import classification_report, confusion_matrix
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Paths
input_path = r"D:\Anik\Jujube\boroi"
output_path = r"D:\Anik\Jujube\Results\MobileNetV2"
os.makedirs(output_path, exist_ok=True)

# Parameters
IMG_SIZE = (224, 224)  # MobileNetV2 preferred input size
BATCH_SIZE = 32
EPOCHS = 25
LEARNING_RATE = 0.0001

# Data Augmentation and Preprocessing
train_datagen = ImageDataGenerator(
    preprocessing_function=preprocess_input,
    validation_split=0.2,
    rotation_range=30,
    width_shift_range=0.2,
    height_shift_range=0.2,
    shear_range=0.2,
    zoom_range=0.2,
    horizontal_flip=True,
    fill_mode="nearest"
)

train_generator = train_datagen.flow_from_directory(
    input_path,
    target_size=IMG_SIZE,
    batch_size=BATCH_SIZE,
    class_mode='categorical',
    subset='training',
    shuffle=True
)

validation_generator = train_datagen.flow_from_directory(
    input_path,
    target_size=IMG_SIZE,
    batch_size=BATCH_SIZE,
    class_mode='categorical',
    subset='validation',
    shuffle=False
)

# Load Pretrained MobileNetV2
base_model = MobileNetV2(weights='imagenet', include_top=False, input_shape=(224, 224, 3))
base_model.trainable = False

# Custom Head
x = base_model.output
x = GlobalAveragePooling2D()(x)
x = Dense(1024, activation='relu')(x)
x = Dropout(0.5)(x)
predictions = Dense(train_generator.num_classes, activation='softmax')(x)

model = Model(inputs=base_model.input, outputs=predictions)

# Compile
model.compile(optimizer=RMSprop(learning_rate=LEARNING_RATE),
              loss='categorical_crossentropy',
              metrics=['accuracy'])

# Callbacks
checkpoint = ModelCheckpoint(os.path.join(output_path, 'best_model.keras'),
                             monitor='val_accuracy',
                             save_best_only=True,
                             mode='max',
                             verbose=1)

early_stop = EarlyStopping(monitor='val_accuracy', patience=5, restore_best_weights=True)

# Train
history = model.fit(
    train_generator,
    epochs=EPOCHS,
    validation_data=validation_generator,
    callbacks=[checkpoint, early_stop]
)

# Plot Accuracy & Loss
plt.figure(figsize=(12, 5))

plt.subplot(1, 2, 1)
plt.plot(history.history['accuracy'], label='Train Accuracy', marker='o')
plt.plot(history.history['val_accuracy'], label='Validation Accuracy', marker='o')
plt.title("Model Accuracy")
plt.xlabel("Epochs")
plt.ylabel("Accuracy")
plt.legend()

plt.subplot(1, 2, 2)
plt.plot(history.history['loss'], label='Train Loss', marker='o')
plt.plot(history.history['val_loss'], label='Validation Loss', marker='o')
plt.title("Model Loss")
plt.xlabel("Epochs")
plt.ylabel("Loss")
plt.legend()

accuracy_loss_curve_path = os.path.join(output_path, 'accuracy_loss_curve.png')
plt.savefig(accuracy_loss_curve_path)
plt.close()

# Evaluate Model
model.load_weights(os.path.join(output_path, 'best_model.keras'))
eval_result = model.evaluate(validation_generator)

# Predictions
Y_pred = model.predict(validation_generator)
y_pred = np.argmax(Y_pred, axis=1)
y_true = validation_generator.classes

report = classification_report(y_true, y_pred, target_names=list(train_generator.class_indices.keys()))

# Confusion Matrix
cm = confusion_matrix(y_true, y_pred)
confusion_matrix_path = os.path.join(output_path, 'confusion_matrix.png')

plt.figure(figsize=(12, 10))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
            xticklabels=train_generator.class_indices.keys(),
            yticklabels=train_generator.class_indices.keys())
plt.xticks(rotation=45)
plt.yticks(rotation=0)
plt.xlabel('Predicted')
plt.ylabel('Actual')
plt.title('Confusion Matrix')
plt.savefig(confusion_matrix_path)
plt.close()

# Generate Word Report
doc = Document()
doc.add_heading('Classification Report - MobileNetV2', 0)

doc.add_heading('Model Performance', level=1)
doc.add_paragraph(f"Final Validation Loss: {eval_result[0]:.4f}")
doc.add_paragraph(f"Final Validation Accuracy: {eval_result[1]:.4f}")

doc.add_heading('Classification Metrics', level=1)
doc.add_paragraph(report)

doc.add_heading('Training History', level=1)
doc.add_picture(accuracy_loss_curve_path, width=Inches(6))

doc.add_heading('Confusion Matrix', level=1)
doc.add_picture(confusion_matrix_path, width=Inches(6))

doc.save(os.path.join(output_path, 'classification_report.docx'))

print("Training with MobileNetV2 completed and results saved.")
