import os
import numpy as np
import tensorflow as tf
import matplotlib.pyplot as plt
from collections import Counter
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications.mobilenet_v2 import MobileNetV2, preprocess_input
from tensorflow.keras.layers import Dense, GlobalAveragePooling2D, Dropout
from tensorflow.keras.models import Model
from tensorflow.keras.optimizers import RMSprop, Adam
from tensorflow.keras.callbacks import ModelCheckpoint, EarlyStopping
from sklearn.utils.class_weight import compute_class_weight
from sklearn.metrics import classification_report, confusion_matrix
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Paths
input_path = r"D:\Research\Jujube\boroi"
output_path = r"D:\Research\Jujube\Results\MobileNetV2"
os.makedirs(output_path, exist_ok=True)

# Parameters
IMG_SIZE = (224, 224)
BATCH_SIZE = 32
INITIAL_EPOCHS = 15
FINE_TUNE_EPOCHS = 10
LEARNING_RATE = 1e-4
FINE_TUNE_LR = 1e-5

# Data Augmentation
train_datagen = ImageDataGenerator(
    preprocessing_function=preprocess_input,
    validation_split=0.2,
    rotation_range=30,
    width_shift_range=0.2,
    height_shift_range=0.2,
    shear_range=0.2,
    zoom_range=0.2,
    horizontal_flip=True,
    fill_mode="nearest"
)

train_generator = train_datagen.flow_from_directory(
    input_path,
    target_size=IMG_SIZE,
    batch_size=BATCH_SIZE,
    class_mode='categorical',
    subset='training',
    shuffle=True
)

validation_generator = train_datagen.flow_from_directory(
    input_path,
    target_size=IMG_SIZE,
    batch_size=BATCH_SIZE,
    class_mode='categorical',
    subset='validation',
    shuffle=False
)

# Compute Class Weights
labels = train_generator.classes
class_weights = compute_class_weight(class_weight='balanced', classes=np.unique(labels), y=labels)
class_weights_dict = dict(enumerate(class_weights))

# Load MobileNetV2
base_model = MobileNetV2(weights='imagenet', include_top=False, input_shape=(224, 224, 3))
base_model.trainable = False

# Head
x = base_model.output
x = GlobalAveragePooling2D()(x)
x = Dense(1024, activation='relu')(x)
x = Dropout(0.5)(x)
predictions = Dense(train_generator.num_classes, activation='softmax')(x)
model = Model(inputs=base_model.input, outputs=predictions)

# Compile initial model
model.compile(optimizer=RMSprop(learning_rate=LEARNING_RATE),
              loss='categorical_crossentropy',
              metrics=['accuracy'])

checkpoint = ModelCheckpoint(os.path.join(output_path, 'best_model.keras'),
                             monitor='val_loss',
                             save_best_only=True,
                             mode='min',
                             verbose=1)

early_stop = EarlyStopping(monitor='val_loss', patience=5, restore_best_weights=True)

# Initial training
history = model.fit(
    train_generator,
    epochs=INITIAL_EPOCHS,
    validation_data=validation_generator,
    callbacks=[checkpoint, early_stop],
    class_weight=class_weights_dict
)

# Fine-tune last few layers
base_model.trainable = True
for layer in base_model.layers[:-40]:
    layer.trainable = False

model.compile(optimizer=Adam(learning_rate=FINE_TUNE_LR),
              loss='categorical_crossentropy',
              metrics=['accuracy'])

fine_tune_history = model.fit(
    train_generator,
    epochs=INITIAL_EPOCHS + FINE_TUNE_EPOCHS,
    initial_epoch=history.epoch[-1] + 1,
    validation_data=validation_generator,
    callbacks=[checkpoint, early_stop],
    class_weight=class_weights_dict
)

# Accuracy & Loss plots
def plot_curves(histories, save_path):
    plt.figure(figsize=(12, 5))
    for history, label in histories:
        plt.plot(history.history['accuracy'], label=f'{label} Train Acc')
        plt.plot(history.history['val_accuracy'], label=f'{label} Val Acc')
    plt.title("Model Accuracy")
    plt.xlabel("Epochs")
    plt.ylabel("Accuracy")
    plt.legend()
    plt.savefig(save_path)
    plt.close()

plot_curves([(history, 'Initial'), (fine_tune_history, 'Fine-Tune')],
            os.path.join(output_path, 'accuracy_loss_curve.png'))

# Evaluate
model.load_weights(os.path.join(output_path, 'best_model.keras'))
eval_result = model.evaluate(validation_generator)

# Classification report
Y_pred = model.predict(validation_generator)
y_pred = np.argmax(Y_pred, axis=1)
y_true = validation_generator.classes
report = classification_report(y_true, y_pred, target_names=list(train_generator.class_indices.keys()))

# Confusion matrix
cm = confusion_matrix(y_true, y_pred)
cm_path = os.path.join(output_path, 'confusion_matrix.png')
plt.figure(figsize=(12, 10))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
            xticklabels=train_generator.class_indices.keys(),
            yticklabels=train_generator.class_indices.keys())
plt.xticks(rotation=45)
plt.yticks(rotation=0)
plt.xlabel('Predicted')
plt.ylabel('Actual')
plt.title('Confusion Matrix')
plt.savefig(cm_path)
plt.close()

# Word report
doc = Document()
doc.add_heading('Classification Report - Improved MobileNetV2', 0)

doc.add_heading('Model Performance', level=1)
doc.add_paragraph(f"Final Validation Loss: {eval_result[0]:.4f}")
doc.add_paragraph(f"Final Validation Accuracy: {eval_result[1]:.4f}")

doc.add_heading('Classification Metrics', level=1)
doc.add_paragraph(report)

doc.add_heading('Training History', level=1)
doc.add_picture(os.path.join(output_path, 'accuracy_loss_curve.png'), width=Inches(6))

doc.add_heading('Confusion Matrix', level=1)
doc.add_picture(cm_path, width=Inches(6))

doc.save(os.path.join(output_path, 'classification_report.docx'))

print("Improved MobileNetV2 training completed and results saved.")
