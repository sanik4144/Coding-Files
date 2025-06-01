# type: ignore

import os
import numpy as np
import matplotlib.pyplot as plt
from tensorflow.keras.applications import VGG19
from tensorflow.keras.layers import Dense, Flatten, Dropout
from tensorflow.keras.models import Model
from tensorflow.keras.optimizers import Adam
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.callbacks import EarlyStopping, ReduceLROnPlateau
from sklearn.metrics import confusion_matrix, precision_score, recall_score, f1_score, classification_report
import seaborn as sns
from docx import Document
from docx.shared import Inches

# === Step 1: Paths & Parameters ===
dataset_path = r"D:\Research\Jujube\boroi"  # UPDATE this with your actual dataset path
output_path = r"D:\Research\Jujube\Results\VGG19"
os.makedirs(output_path, exist_ok=True)

img_size = (224, 224)
batch_size = 32
epochs = 30
learning_rate = 1e-5

# === Step 2: Load Dataset ===
data_gen = ImageDataGenerator(rescale=1.0/255, validation_split=0.2)

train_data = data_gen.flow_from_directory(
    dataset_path,
    target_size=img_size,
    batch_size=batch_size,
    class_mode='binary',
    subset='training',
    shuffle=True
)

val_data = data_gen.flow_from_directory(
    dataset_path,
    target_size=img_size,
    batch_size=batch_size,
    class_mode='binary',
    subset='validation',
    shuffle=False
)

# === Step 3: Build VGG19 Model ===
def build_vgg19():
    base_model = VGG19(weights='imagenet', include_top=False, input_shape=(224, 224, 3))
    
    # Fine-tuning last 4 layers
    for layer in base_model.layers[:-4]:
        layer.trainable = False
    for layer in base_model.layers[-4:]:
        layer.trainable = True

    x = Flatten()(base_model.output)
    x = Dense(512, activation='relu')(x)
    x = Dropout(0.7)(x)
    output = Dense(1, activation='sigmoid')(x)

    model = Model(inputs=base_model.input, outputs=output)
    model.compile(optimizer=Adam(learning_rate=learning_rate),
                  loss='binary_crossentropy',
                  metrics=['accuracy'])
    return model

model = build_vgg19()

# === Step 4: Train Model ===
early_stopping = EarlyStopping(monitor='val_loss', patience=5, restore_best_weights=True)
reduce_lr = ReduceLROnPlateau(monitor='val_loss', factor=0.2, patience=3, min_lr=1e-7, verbose=1)

history = model.fit(
    train_data,
    validation_data=val_data,
    epochs=epochs,
    callbacks=[early_stopping, reduce_lr]
)

# === Step 5: Plot Training History ===
plt.figure(figsize=(12, 5))

# Accuracy
plt.subplot(1, 2, 1)
plt.plot(history.history['accuracy'], label='Train Accuracy')
plt.plot(history.history['val_accuracy'], label='Val Accuracy')
plt.title('Model Accuracy')
plt.xlabel('Epochs')
plt.ylabel('Accuracy')
plt.legend()

# Loss
plt.subplot(1, 2, 2)
plt.plot(history.history['loss'], label='Train Loss')
plt.plot(history.history['val_loss'], label='Val Loss')
plt.title('Model Loss')
plt.xlabel('Epochs')
plt.ylabel('Loss')
plt.legend()

curve_path = os.path.join(output_path, 'accuracy_loss_curve.png')
plt.savefig(curve_path)
plt.close()

# === Step 6: Evaluate Model ===
val_loss, val_accuracy = model.evaluate(val_data)
print(f"Validation Accuracy: {val_accuracy*100:.2f}%")
print(f"Validation Loss: {val_loss:.4f}")

# === Step 7: Predictions & Metrics ===
val_data.reset()
true_labels = val_data.classes
predictions = (model.predict(val_data) > 0.5).astype("int32").flatten()

conf_matrix = confusion_matrix(true_labels, predictions)
precision = precision_score(true_labels, predictions)
recall = recall_score(true_labels, predictions)
f1 = f1_score(true_labels, predictions)
report = classification_report(true_labels, predictions, target_names=['Healthy', 'Not Healthy'])

# === Step 8: Plot Confusion Matrix ===
plt.figure(figsize=(8, 6))
sns.heatmap(conf_matrix, annot=True, fmt='d', cmap='Blues',
            xticklabels=['Healthy', 'Not Healthy'],
            yticklabels=['Healthy', 'Not Healthy'])
plt.title('Confusion Matrix')
plt.xlabel('Predicted')
plt.ylabel('Actual')

conf_matrix_path = os.path.join(output_path, 'confusion_matrix.png')
plt.savefig(conf_matrix_path)
plt.close()

# === Step 9: Save Word Report ===
doc = Document()
doc.add_heading('VGG19 Classification Report', 0)

# Summary
doc.add_heading('Model Performance', level=1)
doc.add_paragraph(f"Validation Accuracy: {val_accuracy*100:.2f}%")
doc.add_paragraph(f"Validation Loss: {val_loss:.4f}")
doc.add_paragraph(f"Precision: {precision:.2f}")
doc.add_paragraph(f"Recall: {recall:.2f}")
doc.add_paragraph(f"F1 Score: {f1:.2f}")

# Classification Report
doc.add_heading('Classification Report', level=1)
doc.add_paragraph(report)

# Accuracy/Loss Curve
doc.add_heading('Training History', level=1)
doc.add_picture(curve_path, width=Inches(6))

# Confusion Matrix
doc.add_heading('Confusion Matrix', level=1)
doc.add_picture(conf_matrix_path, width=Inches(5.5))

# Save Document
doc.save(os.path.join(output_path, 'vgg19_classification_report.docx'))

print("✅ Training completed and results saved!")
